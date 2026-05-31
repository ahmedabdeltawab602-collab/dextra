"""
lessonkit.py — Lesson-level scaffolding.
Produces docx + matching JSON with stable IDs.
"""
import os, json, datetime
from helpers import (setup, titlep, h2, h3, ep, ap, el, eb, ab, en_, an_,
                     eq, div, fc, end_m, callout, kvtable, pgbrk,
                     FONT, PRIMARY, SECONDARY, RED, GREY, YELLOW, GREEN,
                     ORANGE, LIGHT_RED, WHITE, BLACK, DARK_TEXT)

SECTIONS = [
    "intro", "definitions", "concept", "formulas", "example",
    "interpret", "traps", "practice", "summary", "memory", "quiz",
]

class Lesson:
    def __init__(self, *, lid, module, title_en, title_ar,
                 objectives_en=None, objectives_ar=None,
                 sources=None, out_dir=None):
        self.lid = lid
        self.module = module
        self.title_en = title_en
        self.title_ar = title_ar
        self.objectives_en = objectives_en or []
        self.objectives_ar = objectives_ar or []
        self.sources = sources or []
        self.out_dir = out_dir
        self.d = setup()
        fc(self.d)
        self.data = {
            "id": lid, "module": module, "type": "lesson",
            "title_en": title_en, "title_ar": title_ar,
            "objectives_en": self.objectives_en,
            "objectives_ar": self.objectives_ar,
            "sources": self.sources,
            "sections": [], "formulas": [], "examples": [],
            "memory_aid": [], "summary": {"en":"", "ar":""},
            "quiz": [],
            "created": datetime.date.today().isoformat()
        }

    def header(self):
        d = self.d
        from helpers import _run, _HDR
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = d.add_paragraph()
        _HDR(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
        _run(p, f"{self.lid} · Module {self.module:02d}", bold=True, size=11, color=SECONDARY)
        p = d.add_paragraph(); _HDR(p, space_after=2)
        _run(p, self.title_en, bold=True, size=22, color=PRIMARY)
        p = d.add_paragraph(); _HDR(p, rtl=True, space_after=6)
        _run(p, self.title_ar, bold=True, size=20, color=PRIMARY, rtl=True)
        div(d, color=SECONDARY, sz=8, after=6)

    def objectives(self):
        d = self.d
        h2(d, "Learning Objectives", "اهداف الدرس")
        for o in self.objectives_en: eb(d, o)
        for o in self.objectives_ar: ab(d, o)
        div(d, color=GREY, sz=4, after=4)

    def section(self, key, en_title, ar_title, body_callable=None):
        d = self.d
        h2(d, en_title, ar_title)
        if body_callable: body_callable(d)
        self.data["sections"].append({"key":key, "title_en":en_title, "title_ar":ar_title})

    def add_formula(self, *, fid, name_en, name_ar, expr,
                    where_en="", where_ar="", when_en="", when_ar=""):
        self.data["formulas"].append({
            "id": fid, "name_en": name_en, "name_ar": name_ar, "expr": expr,
            "where_en": where_en, "where_ar": where_ar,
            "when_en": when_en, "when_ar": when_ar,
        })
        eq(self.d, expr,
           note_en=f"{name_en}" + (f" — {where_en}" if where_en else ""),
           note_ar=f"{name_ar}" + (f" — {where_ar}" if where_ar else ""))
        if when_en or when_ar:
            callout(self.d, "use",
                    f"Use when: {when_en}" if when_en else "",
                    f"يُستخدم عندما: {when_ar}" if when_ar else "")

    def add_example(self, *, problem_en, problem_ar, steps, answer_en, answer_ar):
        self.data["examples"].append({
            "problem":{"en":problem_en,"ar":problem_ar},
            "steps": steps, "answer":{"en":answer_en,"ar":answer_ar}
        })
        d = self.d
        callout(d, "example", problem_en, problem_ar)
        for i, st in enumerate(steps, 1):
            ep(d, f"Step {i}: {st['en']}")
            ap(d, f"الخطوة {i}: {st['ar']}")
            if 'calc' in st:
                eq(d, st['calc'])
        callout(d, "key", f"Answer: {answer_en}", f"الإجابة: {answer_ar}")

    def add_memory_aid(self, en, ar):
        self.data["memory_aid"].append({"en":en, "ar":ar})

    def summary(self, en, ar):
        self.data["summary"] = {"en":en, "ar":ar}
        h2(self.d, "Summary", "الخلاصة")
        ep(self.d, en); ap(self.d, ar)

    def quiz_q(self, *, qid, qtype, q_en, q_ar, opts_en=None, opts_ar=None,
               ans, exp_en, exp_ar, formula_used=""):
        self.data["quiz"].append({
            "id": qid, "topic": self.lid, "type": qtype,
            "q_en": q_en, "q_ar": q_ar,
            "opts_en": opts_en or [], "opts_ar": opts_ar or [],
            "ans": ans, "exp_en": exp_en, "exp_ar": exp_ar,
            "formula_used": formula_used
        })

    def render_quiz(self):
        if not self.data["quiz"]: return
        d = self.d
        h2(d, "Quiz - 5 Questions", "كويز - 5 اسئلة")
        from helpers import _run, _HDR
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for i, q in enumerate(self.data["quiz"], 1):
            p = d.add_paragraph(); _HDR(p, space_before=4, space_after=1)
            _run(p, f"Q{i}. ", bold=True, size=12, color=SECONDARY)
            _run(p, q["q_en"], bold=False, size=11)
            p2 = d.add_paragraph(); _HDR(p2, rtl=True, space_after=2)
            _run(p2, f"س{i}. ", bold=True, size=12, color=SECONDARY, rtl=True)
            _run(p2, q["q_ar"], size=11, rtl=True)
            if q["opts_en"]:
                for idx, opt in enumerate(q["opts_en"]):
                    L = chr(ord('A')+idx)
                    eb(d, f"{L}. {opt}")
                for idx, opt in enumerate(q["opts_ar"]):
                    L = chr(ord('A')+idx)
                    ab(d, f"{L}. {opt}")
            callout(d, "key",
                    f"Answer: {q['ans']} - {q['exp_en']}",
                    f"الإجابة: {q['ans']} - {q['exp_ar']}")
            if q.get("formula_used"):
                ep(d, f"Formula used: {q['formula_used']}", italic=True, color=SECONDARY)

    def render_memory(self):
        if not self.data["memory_aid"]: return
        d = self.d
        h2(d, "Memory Aids - Reference Card", "بطاقة القوانين")
        for m in self.data["memory_aid"]:
            callout(d, "tip", m["en"], m["ar"])

    def save(self):
        os.makedirs(self.out_dir, exist_ok=True)
        # Sanitize filename for Windows: ASCII only dashes, no special chars
        t = self.title_en
        for bad, good in [("/", "-"), (":", " -"), ("&", "and"),
                          ("—", "-"), ("–", "-"),
                          ("?", ""), ("*", ""), ('"', ""),
                          ("<", ""), (">", ""), ("|", "")]:
            t = t.replace(bad, good)
        safe = t[:80].strip()
        base = f"{self.lid} - {safe}"
        docx_path = os.path.join(self.out_dir, base + ".docx")
        json_path = os.path.join(self.out_dir, base + ".json")
        end_m(self.d, docx_path)
        with open(json_path, "w", encoding="utf-8") as fp:
            json.dump(self.data, fp, ensure_ascii=False, indent=2)
        return docx_path, json_path
