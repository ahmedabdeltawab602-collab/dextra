"""
helpers.py — Statistics Course Build
Compact docx helpers + RTL/LTR bilingual paragraphs + styled blocks.
Author: Ahmed Abd El Tawab
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy, re, json, os

# ─────────── Design constants ───────────
FONT       = "Arial"
PRIMARY    = "1F4E79"
SECONDARY  = "2E75B6"
RED        = "C00000"
GREY       = "F2F2F2"
YELLOW     = "FFF2CC"
GREEN      = "C6E0B4"
ORANGE     = "F4B084"
LIGHT_RED  = "F8CBAD"
WHITE      = "FFFFFF"
BLACK      = "000000"
DARK_TEXT  = "1F1F1F"

# ─────────── Internal helpers ───────────
def _hex(c): return c.lstrip("#")

def _run(p, text, *, bold=False, italic=False, size=11, color=DARK_TEXT, font=FONT, rtl=False):
    """Add a styled run to paragraph p."""
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    rf = r.font
    rf.name = font
    rf.size = Pt(size)
    rf.color.rgb = RGBColor.from_string(_hex(color))
    rPr = r._element.get_or_add_rPr()
    # ASCII / HAnsi / cs / EastAsia fonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
        rFonts.set(qn(a), font)
    if rtl:
        rtl_el = OxmlElement('w:rtl'); rtl_el.set(qn('w:val'), 'true')
        rPr.append(rtl_el)
        # complex-script size
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(size*2)))
        rPr.append(szCs)
    return r

def _HDR(p, *, rtl=False, align=None, space_before=0, space_after=4, keep_with_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if keep_with_next: pf.keep_with_next = True
    pPr = p._element.get_or_add_pPr()
    if rtl:
        bd = OxmlElement('w:bidi'); bd.set(qn('w:val'), 'true'); pPr.append(bd)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align is None else align
    else:
        # ensure LTR
        bd = pPr.find(qn('w:bidi'))
        if bd is not None: pPr.remove(bd)
        if align is not None: p.alignment = align
    return p

def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _hex(fill_hex))
    tcPr.append(shd)

def _set_cell_borders(cell, color=PRIMARY, sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz)); b.set(qn('w:color'),_hex(color))
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _normalize(doc):
    """Strip null bytes / non-XML chars from text runs."""
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text and ('\x00' in r.text):
                r.text = r.text.replace('\x00','')
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text and ('\x00' in r.text):
                            r.text = r.text.replace('\x00','')

# ─────────── Public compact API ───────────
def setup(margins_cm=2.0):
    """Create a new Document with 2cm margins."""
    d = Document()
    # default style font
    style = d.styles['Normal']
    style.font.name = FONT; style.font.size = Pt(11)
    # margins
    for section in d.sections:
        section.top_margin    = Cm(margins_cm)
        section.bottom_margin = Cm(margins_cm)
        section.left_margin   = Cm(margins_cm)
        section.right_margin  = Cm(margins_cm)
    return d

def titlep(d, *, module_no, module_en, module_ar, subtitle_en="", subtitle_ar=""):
    """Title page block (no separate page break — caller decides)."""
    # spacer
    sp = d.add_paragraph(); _HDR(sp); _run(sp,"",size=11)
    sp.paragraph_format.space_before = Pt(60)
    # Module number
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, f"MODULE {module_no:02d}", bold=True, size=20, color=SECONDARY)
    # English title
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _run(p, module_en, bold=True, size=28, color=PRIMARY)
    # Arabic title
    p = d.add_paragraph(); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    _run(p, module_ar, bold=True, size=24, color=PRIMARY, rtl=True)
    if subtitle_en:
        p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        _run(p, subtitle_en, italic=True, size=13, color=DARK_TEXT)
    if subtitle_ar:
        p = d.add_paragraph(); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        _run(p, subtitle_ar, italic=True, size=13, color=DARK_TEXT, rtl=True)
    # divider line
    div(d, color=SECONDARY)
    # author
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
    _run(p, "Ahmed Abd El Tawab — Statistics Course", italic=True, size=11, color=SECONDARY)

def h2(d, en, ar=None, *, color=PRIMARY):
    """Section heading (English + Arabic on next line)."""
    p = d.add_paragraph(); _HDR(p, space_before=10, space_after=2, keep_with_next=True)
    _run(p, en, bold=True, size=16, color=color)
    if ar:
        p2 = d.add_paragraph(); _HDR(p2, rtl=True, space_before=0, space_after=4, keep_with_next=True)
        _run(p2, ar, bold=True, size=14, color=color, rtl=True)
    # underline bar
    div(d, color=color, sz=8, after=4)

def h3(d, en, ar=None, *, color=SECONDARY):
    """Sub-section heading."""
    p = d.add_paragraph(); _HDR(p, space_before=6, space_after=1, keep_with_next=True)
    _run(p, en, bold=True, size=13, color=color)
    if ar:
        p2 = d.add_paragraph(); _HDR(p2, rtl=True, space_after=2, keep_with_next=True)
        _run(p2, ar, bold=True, size=12, color=color, rtl=True)

def ep(d, text, *, size=11, color=DARK_TEXT, bold=False, italic=False):
    """English paragraph (LTR, justified)."""
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _run(p, text, size=size, color=color, bold=bold, italic=italic)
    return p

def ap(d, text, *, size=11, color=DARK_TEXT, bold=False, italic=False):
    """Arabic paragraph (RTL, justified)."""
    p = d.add_paragraph(); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, text, size=size, color=color, bold=bold, italic=italic, rtl=True)
    return p

def el(d, en, ar=None):
    """Bilingual one-line block — En then Ar."""
    ep(d, en)
    if ar: ap(d, ar)

def al(d, ar):
    """Arabic-only line."""
    ap(d, ar)

def eb(d, text, *, size=11):
    """English bullet."""
    p = d.add_paragraph(style='List Bullet'); _HDR(p, space_after=1)
    _run(p, text, size=size)
    return p

def ab(d, text, *, size=11):
    """Arabic bullet (RTL)."""
    p = d.add_paragraph(style='List Bullet'); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=3)
    _run(p, text, size=size, rtl=True)
    return p

def en_(d, text, *, size=11):
    """English numbered item."""
    p = d.add_paragraph(style='List Number'); _HDR(p, space_after=1)
    _run(p, text, size=size)
    return p

def an_(d, text, *, size=11):
    """Arabic numbered item (RTL)."""
    p = d.add_paragraph(style='List Number'); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=3)
    _run(p, text, size=size, rtl=True)
    return p

def eq(d, expr, *, note_en=None, note_ar=None):
    """Equation in Unicode — centered, monospace-ish, with light shading."""
    tbl = d.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    cell = tbl.cell(0,0)
    _shade(cell, GREY); _set_cell_borders(cell, color=SECONDARY, sz=4)
    p = cell.paragraphs[0]; _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    _run(p, expr, bold=True, size=14, color=PRIMARY, font="Cambria Math")
    if note_en: ep(d, note_en, italic=True, color=SECONDARY)
    if note_ar: ap(d, note_ar, italic=True, color=SECONDARY)

def div(d, *, color=PRIMARY, sz=6, after=4):
    """Horizontal divider line."""
    p = d.add_paragraph(); _HDR(p, space_after=after)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),_hex(color))
    pBdr.append(bot); pPr.append(pBdr)

def fc(d, footer_text="Ahmed Abd El Tawab — Statistics Course"):
    """Footer with page number."""
    section = d.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"{footer_text} — Page ", size=10, color=SECONDARY, italic=True)
    # PAGE field
    r = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r._element.append(fld)
    rf = r.font; rf.name = FONT; rf.size = Pt(10)
    rf.color.rgb = RGBColor.from_string(_hex(SECONDARY))

def end_m(d, path, *, lang_pair=True):
    """Final save with OOXML normalize."""
    _normalize(d)
    d.save(path)

# ─────────── Special blocks ───────────
def callout(d, kind, en, ar=None):
    """Colored callout: kind ∈ {tip, warning, key, example, trap}"""
    palette = {
        'tip':     (GREEN,     "💡 TIP",     "💡 نصيحة"),
        'warning': (LIGHT_RED, "⚠ WARNING", "⚠ تحذير"),
        'key':     (YELLOW,    "★ KEY IDEA","★ فكرة محورية"),
        'example': (GREY,      "🧮 EXAMPLE","🧮 مثال"),
        'trap':    (ORANGE,    "🪤 COMMON TRAP","🪤 فخ شائع"),
        'meaning': (YELLOW,    "🎯 MEANING","🎯 الدلالة"),
        'use':     (GREEN,     "🛠 IN PRACTICE","🛠 الاستخدام العملى"),
    }
    fill, lbl_en, lbl_ar = palette.get(kind, (GREY, "NOTE", "ملاحظة"))
    tbl = d.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0,0); _shade(cell, fill); _set_cell_borders(cell, color=SECONDARY, sz=4)
    # label line
    lp = cell.paragraphs[0]; _HDR(lp, space_after=2)
    _run(lp, lbl_en, bold=True, size=11, color=PRIMARY)
    # english body
    pe = cell.add_paragraph(); _HDR(pe, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _run(pe, en, size=11)
    if ar:
        # arabic label
        pa_l = cell.add_paragraph(); _HDR(pa_l, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
        _run(pa_l, lbl_ar, bold=True, size=11, color=PRIMARY, rtl=True)
        pa = cell.add_paragraph(); _HDR(pa, rtl=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
        _run(pa, ar, size=11, rtl=True)
    # spacing after the block
    sp = d.add_paragraph(); _HDR(sp, space_after=4)

def kvtable(d, rows, *, header=None, col_widths=None):
    """Two/three-column comparison table. rows = list of tuples."""
    ncols = len(rows[0]) if rows else 2
    tbl = d.add_table(rows=0, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        hr = tbl.add_row()
        for i, h in enumerate(header):
            c = hr.cells[i]; _shade(c, PRIMARY); _set_cell_borders(c, color=PRIMARY, sz=8)
            p = c.paragraphs[0]; _HDR(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            _run(p, h, bold=True, size=11, color=WHITE)
    for ri, row in enumerate(rows):
        rr = tbl.add_row()
        fill = WHITE if ri % 2 == 0 else GREY
        for i, val in enumerate(row):
            c = rr.cells[i]; _shade(c, fill); _set_cell_borders(c, color=SECONDARY, sz=4)
            p = c.paragraphs[0]; _HDR(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
            # auto-detect Arabic
            is_ar = bool(re.search(r'[؀-ۿ]', val))
            if is_ar:
                _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                _run(p, val, size=11, rtl=True)
            else:
                _run(p, val, size=11)

def page_break(d):
    d.add_paragraph().add_run().add_break(6)  # WD_BREAK.PAGE = 7 ; use add_page_break shortcut
def pgbrk(d):
    d.add_page_break()
