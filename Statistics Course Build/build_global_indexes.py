"""Build global indexes:
   - index.json     → master catalog of every module, lesson, section, formula, example, quiz.
   - formulas.json  → focused cross-module formula registry for the future app.
   - course-stats.json → quick stats (counts) for dashboards.
   Also produces a printable 'Course Index.docx'.
"""
import sys, os, json, datetime
sys.path.insert(0, "/sessions/bold-great-lamport/mnt/02 Statistics Fundamentals/Statistics Course Build")
from helpers import (h2, h3, ep, ap, el, eb, ab, en_, an_, eq, div, callout, kvtable, pgbrk,
                     setup, titlep, fc, end_m, _run, _HDR,
                     PRIMARY, SECONDARY, RED, GREY, YELLOW, GREEN, ORANGE, LIGHT_RED, WHITE)
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "/sessions/bold-great-lamport/mnt/02 Statistics Fundamentals/Statistics Course Build"

# ───────── Walk every module → every lesson JSON ─────────
modules_meta = [
    ("M01", "Module 01 - Introduction",                "Introduction",                  "المقدمة"),
    ("M02", "Module 02 - The Statistical Landscape",   "The Statistical Landscape",     "المشهد الإحصائى"),
    ("M03", "Module 03 - Descriptive Statistics",      "Descriptive Statistics",        "الإحصاء الوصفى"),
    ("M04", "Module 04 - Inferential Statistics",      "Inferential Statistics",        "الإحصاء الاستدلالى"),
    ("M05", "Module 05 - Hypothesis Testing",          "Hypothesis Testing",            "اختبار الفرضيات"),
    ("M06", "Module 06 - Conclusion",                  "Conclusion",                    "الخلاصة"),
    ("M07", "Module 07 - Question Bank",               "Question Bank",                 "بنك الأسئلة"),
]

global_index = {
    "course_id": "STATS-FUND",
    "course_title_en": "Statistics Fundamentals",
    "course_title_ar": "أساسيات الإحصاء",
    "author": "Ahmed Abd El Tawab",
    "generated": datetime.date.today().isoformat(),
    "modules": []
}
formulas_registry = []   # flat list across modules
quiz_count = 0
lesson_count = 0
section_count = 0
example_count = 0
memory_aid_count = 0

for mid, folder, en, ar in modules_meta:
    mdir = os.path.join(ROOT, folder)
    if not os.path.isdir(mdir): continue
    module_entry = {
        "id": mid, "title_en": en, "title_ar": ar,
        "folder": folder, "lessons": []
    }
    files = sorted(f for f in os.listdir(mdir) if f.endswith(".json"))
    for fn in files:
        with open(os.path.join(mdir, fn), encoding="utf-8") as fp:
            data = json.load(fp)
        # skip module index files
        if data.get("type") in ("module", "question_bank"): continue
        lesson_count += 1
        section_count += len(data.get("sections", []))
        example_count += len(data.get("examples", []))
        memory_aid_count += len(data.get("memory_aid", []))
        quiz_count += len(data.get("quiz", []))

        lesson_entry = {
            "id": data["id"],
            "title_en": data["title_en"],
            "title_ar": data["title_ar"],
            "filename_docx": fn.replace(".json", ".docx"),
            "filename_json": fn,
            "sections":     [s["key"] for s in data.get("sections", [])],
            "n_formulas":   len(data.get("formulas", [])),
            "n_examples":   len(data.get("examples", [])),
            "n_memory":     len(data.get("memory_aid", [])),
            "n_quiz":       len(data.get("quiz", [])),
            "sources":      data.get("sources", [])
        }
        module_entry["lessons"].append(lesson_entry)
        # Pull formulas
        for f in data.get("formulas", []):
            formulas_registry.append({
                **f,
                "module_id": mid,
                "lesson_id": data["id"],
                "lesson_title_en": data["title_en"],
                "lesson_title_ar": data["title_ar"]
            })
    global_index["modules"].append(module_entry)

global_index["stats"] = {
    "modules": len([m for m in global_index["modules"] if m["lessons"]]),
    "lessons": lesson_count,
    "sections": section_count,
    "examples": example_count,
    "memory_aids": memory_aid_count,
    "quiz_questions": quiz_count,
    "formulas": len(formulas_registry),
}

# ───────── Save JSON indexes ─────────
with open(os.path.join(ROOT, "index.json"), "w", encoding="utf-8") as fp:
    json.dump(global_index, fp, ensure_ascii=False, indent=2)
print("index.json written.")

formulas_doc = {
    "course_id": "STATS-FUND",
    "title_en": "Statistics Fundamentals - Formula Registry",
    "title_ar": "أساسيات الإحصاء - سجل القوانين",
    "generated": datetime.date.today().isoformat(),
    "total_formulas": len(formulas_registry),
    "formulas": formulas_registry
}
with open(os.path.join(ROOT, "formulas.json"), "w", encoding="utf-8") as fp:
    json.dump(formulas_doc, fp, ensure_ascii=False, indent=2)
print(f"formulas.json written ({len(formulas_registry)} formulas).")

with open(os.path.join(ROOT, "course-stats.json"), "w", encoding="utf-8") as fp:
    json.dump(global_index["stats"], fp, indent=2)
print("course-stats.json written.")

# ───────── Build printable Course Index.docx ─────────
d = setup(); fc(d, footer_text="Ahmed Abd El Tawab - Course Index")
titlep(d, module_no=0,
       module_en="Course Index & Formula Registry",
       module_ar="فهرس الدورة وسجل القوانين",
       subtitle_en="The complete map of the Statistics Fundamentals course",
       subtitle_ar="الخريطة الكاملة لدورة أساسيات الإحصاء")

# Stats summary
pgbrk(d)
h2(d, "Course at a Glance", "نظرة عامة على الدورة")
s = global_index["stats"]
kvtable(d, header=["Metric / المقياس","Count / العدد"],
        rows=[
            ("Modules (substantive)\nالموديولات", str(s["modules"])),
            ("Lessons\nالدروس", str(s["lessons"])),
            ("Sections (lesson chapters)\nالأقسام", str(s["sections"])),
            ("Worked examples\nأمثلة محلولة", str(s["examples"])),
            ("Memory aids / reference cards\nبطاقات قوانين", str(s["memory_aids"])),
            ("Quiz questions (bilingual)\nأسئلة كويز", str(s["quiz_questions"])),
            ("Distinct formulas\nقوانين مميَّزة", str(s["formulas"])),
        ])

# Per-module lesson tree
for m in global_index["modules"]:
    if not m["lessons"]: continue
    pgbrk(d)
    h2(d, f"{m['id']} - {m['title_en']}", f"{m['id']} - {m['title_ar']}")
    rows = []
    for L in m["lessons"]:
        rows.append((
            L["id"],
            f"{L['title_en']}\n{L['title_ar']}",
            f"{L['n_quiz']} Q · {L['n_formulas']} F · {L['n_examples']} E"
        ))
    kvtable(d, header=["ID / المعرف","Title / العنوان","Counts / الأعداد"], rows=rows)

# Formula registry
pgbrk(d)
h2(d, "Formula Registry - All Formulas Across the Course",
     "سجل القوانين - كل القوانين عبر الدورة")
ep(d, f"All {len(formulas_registry)} distinct formulas, listed module-by-module.")
ap(d, f"كل القوانين الـ {len(formulas_registry)} مرتبة حسب الموديول.")

# Group formulas by module
by_module = {}
for f in formulas_registry:
    by_module.setdefault(f["module_id"], []).append(f)

for mid, _, en, ar in modules_meta:
    fs = by_module.get(mid, [])
    if not fs: continue
    h3(d, f"{mid} - {en}", f"{mid} - {ar}")
    for f in fs:
        # Formula header
        p = d.add_paragraph(); _HDR(p, space_before=6, space_after=1)
        _run(p, f"[{f['id']}] ", bold=True, size=10, color=SECONDARY)
        _run(p, f["name_en"], bold=True, size=11, color=PRIMARY)
        _run(p, "  ·  ", size=10, color=GREY)
        _run(p, f["name_ar"], bold=True, size=11, color=PRIMARY, rtl=False)
        # Expression
        eq(d, f["expr"])
        # Where & when
        if f.get("where_en") or f.get("where_ar"):
            p = d.add_paragraph(); _HDR(p, space_after=1)
            _run(p, f"   where: {f.get('where_en','')}", italic=True, size=10, color=SECONDARY)
            if f.get("where_ar"):
                p2 = d.add_paragraph(); _HDR(p2, rtl=True, space_after=1)
                _run(p2, f"حيث: {f['where_ar']}", italic=True, size=10, color=SECONDARY, rtl=True)
        if f.get("when_en") or f.get("when_ar"):
            p = d.add_paragraph(); _HDR(p, space_after=4)
            _run(p, f"   use when: {f.get('when_en','')}", italic=True, size=10, color=GREEN)
            if f.get("when_ar"):
                p2 = d.add_paragraph(); _HDR(p2, rtl=True, space_after=4)
                _run(p2, f"يُستخدم عندما: {f['when_ar']}", italic=True, size=10, color=GREEN, rtl=True)
        # tiny separator
        p = d.add_paragraph(); _HDR(p, space_after=0)
        _run(p, "─" * 80, size=8, color=GREY)

end_m(d, os.path.join(ROOT, "Course Index.docx"))
print("Course Index.docx written.")

# Plain-text formula table (one-line / formula) for quick browsing
with open(os.path.join(ROOT, "formulas-quickref.txt"), "w", encoding="utf-8") as fp:
    fp.write("STATISTICS FUNDAMENTALS - FORMULA QUICK REFERENCE\n")
    fp.write("=" * 60 + "\n\n")
    for f in formulas_registry:
        fp.write(f"[{f['id']:18s}]  {f['name_en']:<45s}\n")
        fp.write(f"                    Lesson: {f['lesson_id']}\n")
        fp.write(f"                    Expr:   {f['expr']}\n")
        if f.get("when_en"):
            fp.write(f"                    Use:    {f['when_en']}\n")
        fp.write("\n")
print("formulas-quickref.txt written.")
