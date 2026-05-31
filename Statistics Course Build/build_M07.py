"""Build Module 07 — Comprehensive Question Bank.
Aggregates all quizzes from M01–M06 into:
  - Question Bank (questions only, organised by module)
  - Answer Key (answers + explanations)
  - Master JSON for the future app
"""
import sys, os, json
sys.path.insert(0, "/sessions/bold-great-lamport/mnt/02 Statistics Fundamentals/Statistics Course Build")
from helpers import (h2, h3, ep, ap, el, eb, ab, en_, an_, eq, div, callout, kvtable, pgbrk,
                     setup, titlep, fc, end_m, _run, _HDR,
                     PRIMARY, SECONDARY, RED, GREY, YELLOW, GREEN, ORANGE, LIGHT_RED, WHITE)
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "/sessions/bold-great-lamport/mnt/02 Statistics Fundamentals/Statistics Course Build"
OUT  = os.path.join(ROOT, "Module 07 - Question Bank")
os.makedirs(OUT, exist_ok=True)

# ───────── 1) Collect all quiz items ─────────
modules = [
    ("M01", "Module 01 - Introduction",                     "Introduction",                   "المقدمة"),
    ("M02", "Module 02 - The Statistical Landscape",        "The Statistical Landscape",       "المشهد الإحصائى"),
    ("M03", "Module 03 - Descriptive Statistics",           "Descriptive Statistics",          "الإحصاء الوصفى"),
    ("M04", "Module 04 - Inferential Statistics",           "Inferential Statistics",          "الإحصاء الاستدلالى"),
    ("M05", "Module 05 - Hypothesis Testing",               "Hypothesis Testing",              "اختبار الفرضيات"),
    ("M06", "Module 06 - Conclusion",                       "Conclusion - Review",             "خلاصة - مراجعة"),
]

all_items = []   # flat list with global_index
for mid, folder, en, ar in modules:
    mod_dir = os.path.join(ROOT, folder)
    if not os.path.isdir(mod_dir):
        continue
    files = sorted(f for f in os.listdir(mod_dir)
                   if f.endswith(".json") and "Module Index" not in f)
    for fn in files:
        with open(os.path.join(mod_dir, fn), encoding="utf-8") as fp:
            data = json.load(fp)
        for q in data.get("quiz", []):
            all_items.append({
                "module_id": mid,
                "module_en": en, "module_ar": ar,
                "lesson_id": data["id"],
                "lesson_title_en": data["title_en"],
                "lesson_title_ar": data["title_ar"],
                "qid": q["id"],
                "qtype": q["type"],
                "q_en": q["q_en"], "q_ar": q["q_ar"],
                "opts_en": q.get("opts_en", []),
                "opts_ar": q.get("opts_ar", []),
                "ans": q["ans"],
                "exp_en": q["exp_en"], "exp_ar": q["exp_ar"],
                "formula_used": q.get("formula_used", "")
            })
for i, it in enumerate(all_items, 1):
    it["global_index"] = i

print(f"Collected {len(all_items)} questions.")

# ───────── 2) Save master JSON ─────────
bank_json = {
    "id": "M07",
    "type": "question_bank",
    "title_en": "Statistics Fundamentals - Comprehensive Question Bank",
    "title_ar": "أساسيات الإحصاء - بنك الأسئلة الشامل",
    "modules_covered": [m[0] for m in modules],
    "total_questions": len(all_items),
    "questions": all_items
}
with open(os.path.join(OUT, "M07 - Question Bank Master.json"), "w", encoding="utf-8") as fp:
    json.dump(bank_json, fp, ensure_ascii=False, indent=2)

# ───────── 3) Build Question Bank DOCX (questions only) ─────────
d = setup(); fc(d, footer_text="Ahmed Abd El Tawab - Question Bank")
titlep(d, module_no=7,
       module_en="Comprehensive Question Bank",
       module_ar="بنك الأسئلة الشامل",
       subtitle_en=f"{len(all_items)} bilingual questions across Modules 01-06",
       subtitle_ar=f"{len(all_items)} سؤالاً ثنائى اللغة عبر الموديولات 01-06")

# ToC summary
pgbrk(d)
h2(d, "Contents", "المحتويات")
counts = {}
for it in all_items: counts[it["module_id"]] = counts.get(it["module_id"], 0) + 1
kvtable(d, header=["Module / الموديول","Title / العنوان","# Questions / عدد الأسئلة"],
        rows=[(m[0], f"{m[2]}\n{m[3]}", str(counts.get(m[0],0))) for m in modules])

# Helper to render one question (no answer)
def render_question(d, it):
    # Question header bar
    p = d.add_paragraph()
    _HDR(p, space_before=8, space_after=2)
    _run(p, f"Q{it['global_index']:03d}. ", bold=True, size=12, color=PRIMARY)
    _run(p, f"[{it['lesson_id']} · {it['qtype']}]", size=10, color=SECONDARY, italic=True)

    # English question
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _run(p, it["q_en"], size=11)

    # Arabic question
    p = d.add_paragraph(); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    _run(p, it["q_ar"], size=11, rtl=True)

    # Options
    if it["opts_en"]:
        for idx, opt in enumerate(it["opts_en"]):
            letter = chr(ord('A')+idx)
            p = d.add_paragraph(); _HDR(p, space_after=0)
            _run(p, f"   {letter}. ", bold=True, size=11, color=SECONDARY)
            _run(p, opt, size=11)
        for idx, opt in enumerate(it["opts_ar"]):
            letter = chr(ord('A')+idx)
            p = d.add_paragraph(); _HDR(p, rtl=True, space_after=0)
            _run(p, f"{letter}. ", bold=True, size=11, color=SECONDARY, rtl=True)
            _run(p, opt, size=11, rtl=True)
    # answer line placeholder
    p = d.add_paragraph(); _HDR(p, space_before=4, space_after=2)
    _run(p, "  Answer: __________ ", italic=True, size=10, color=GREY)

# Render module-by-module
for mid, folder, en, ar in modules:
    items = [it for it in all_items if it["module_id"] == mid]
    if not items: continue
    pgbrk(d)
    h2(d, f"{mid} - {en}", f"{mid} - {ar}")
    ep(d, f"{len(items)} questions in this module.")
    ap(d, f"{len(items)} سؤالاً فى هذا الموديول.")
    div(d, color=PRIMARY, sz=6, after=4)
    for it in items:
        render_question(d, it)

end_m(d, os.path.join(OUT, "M07 - Question Bank (Questions Only).docx"))
print("Question Bank docx written.")

# ───────── 4) Build Answer Key DOCX ─────────
d = setup(); fc(d, footer_text="Ahmed Abd El Tawab - Answer Key")
titlep(d, module_no=7,
       module_en="Answer Key",
       module_ar="مفتاح الإجابات",
       subtitle_en=f"Full answers and explanations for all {len(all_items)} questions",
       subtitle_ar=f"إجابات كاملة وشرح لكل الـ {len(all_items)} سؤالاً")

# Quick lookup table (compact)
pgbrk(d)
h2(d, "Quick Answer Lookup", "جدول الإجابات السريع")
ep(d, "Compact answers — full explanations follow by module.")
ap(d, "إجابات مختصرة — شرح كامل يأتى لاحقاً.")

# Build compact 5-column lookup (Q# + Answer)
cols = 4
rows_data = []
for it in all_items:
    rows_data.append((f"Q{it['global_index']:03d}", it["ans"]))
# Render as table with 4 pairs per row = 8 cols
from docx.enum.table import WD_TABLE_ALIGNMENT
chunks = [rows_data[i:i+cols] for i in range(0, len(rows_data), cols)]
tbl = d.add_table(rows=0, cols=cols*2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
from helpers import _shade, _set_cell_borders
for chunk in chunks:
    row = tbl.add_row()
    for i, (qno, ans) in enumerate(chunk):
        c1 = row.cells[i*2]; c2 = row.cells[i*2+1]
        _shade(c1, GREY); _set_cell_borders(c1, color=SECONDARY, sz=4)
        _shade(c2, WHITE); _set_cell_borders(c2, color=SECONDARY, sz=4)
        p1 = c1.paragraphs[0]; _HDR(p1, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p1, qno, bold=True, size=10, color=PRIMARY)
        p2 = c2.paragraphs[0]; _HDR(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p2, ans, bold=True, size=11, color=RED)
    # pad row if uneven
    while len(chunk) < cols:
        i = len(chunk)
        chunk = list(chunk) + [("", "")]
        row.cells[i*2]; row.cells[i*2+1]

# Full explanations module-by-module
def render_answer(d, it):
    # header
    p = d.add_paragraph(); _HDR(p, space_before=6, space_after=2)
    _run(p, f"Q{it['global_index']:03d}. ", bold=True, size=12, color=PRIMARY)
    _run(p, f"[{it['lesson_id']}] Answer: ", bold=True, size=11, color=SECONDARY)
    _run(p, it["ans"], bold=True, size=12, color=RED)
    # English explanation
    p = d.add_paragraph(); _HDR(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _run(p, f"   {it['exp_en']}", size=10)
    # Arabic explanation
    p = d.add_paragraph(); _HDR(p, rtl=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    _run(p, it["exp_ar"], size=10, rtl=True)
    if it.get("formula_used"):
        p = d.add_paragraph(); _HDR(p, space_after=4)
        _run(p, f"   Formula: {it['formula_used']}", italic=True, size=10, color=SECONDARY)

for mid, folder, en, ar in modules:
    items = [it for it in all_items if it["module_id"] == mid]
    if not items: continue
    pgbrk(d)
    h2(d, f"{mid} - {en}", f"{mid} - {ar}")
    div(d, color=PRIMARY, sz=6, after=4)
    for it in items:
        render_answer(d, it)

end_m(d, os.path.join(OUT, "M07 - Answer Key.docx"))
print("Answer Key docx written.")

# ───────── 5) Module index ─────────
module_index = {
    "id": "M07", "module": 7, "type": "module",
    "title_en": "Comprehensive Question Bank",
    "title_ar": "بنك الأسئلة الشامل",
    "files": [
        "M07 - Cover.docx",
        "M07 - Question Bank (Questions Only).docx",
        "M07 - Answer Key.docx",
        "M07 - Question Bank Master.json",
    ],
    "total_questions": len(all_items),
    "breakdown": {m[0]: counts.get(m[0], 0) for m in modules}
}
with open(os.path.join(OUT, "M07 - Module Index.json"), "w", encoding="utf-8") as fp:
    json.dump(module_index, fp, ensure_ascii=False, indent=2)

# Cover docx
d = setup(); fc(d)
titlep(d, module_no=7,
       module_en="Question Bank",
       module_ar="بنك الأسئلة",
       subtitle_en=f"{len(all_items)} questions · 6 modules · bilingual",
       subtitle_ar=f"{len(all_items)} سؤالاً · 6 موديولات · ثنائى اللغة")
end_m(d, os.path.join(OUT, "M07 - Cover.docx"))

print("Module 07 cover and index written.")
print(f"Total: {len(all_items)} questions across {len([m for m in modules if counts.get(m[0],0)])} modules.")
